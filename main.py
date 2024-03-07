from docx import Document
import tkinter as tk
from tkinter import filedialog

from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_page_layout(doc):
    for section in doc.sections:
        # A4
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Mm(30)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(20)
        section.gutter = Mm(10)
        section.footer_distance = Cm(1)
        section.header_distance = Cm(2)


def select_word_document():
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(title="选择一个文档", filetypes=[("Word", "*.docx")])
    return file_path


def custom_save_dialog(original_path):
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", initialdir=original_path.rsplit('/', 1)[0],
                                             title="另存为", filetypes=[("Word", "*.docx")])
    return file_path


def update_headers_if_text_exists(doc, header_text):
    for section in doc.sections:
        if any(paragraph.text.strip() for paragraph in section.header.paragraphs):
            clear_and_set_new_header(section.header, header_text)
        if not section.first_page_header.is_linked_to_previous:
            if any(paragraph.text.strip() for paragraph in section.first_page_header.paragraphs):
                clear_and_set_new_header(section.first_page_header, header_text)
        if section.even_page_header and not section.even_page_header.is_linked_to_previous:
            if any(paragraph.text.strip() for paragraph in section.even_page_header.paragraphs):
                clear_and_set_new_header(section.even_page_header, header_text)


def clear_and_set_new_header(header, text):
    for paragraph in header.paragraphs:
        paragraph.clear()
    new_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    run = new_paragraph.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


if __name__ == "__main__":
    selected_doc_path = select_word_document()
    if selected_doc_path:
        doc = Document(selected_doc_path)
        set_page_layout(doc)
        header_text = "杭州电子科技大学信息工程学院本科毕业设计"
        update_headers_if_text_exists(doc, header_text)

        new_doc_path = custom_save_dialog(selected_doc_path)
        if new_doc_path:
            doc.save(new_doc_path)
            print(f"文件另存为 {new_doc_path}")
        else:
            print("取消保存文件")
    else:
        print("未选择文件或者取消")
